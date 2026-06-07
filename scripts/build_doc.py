#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""出题 Word 生成器：读 problems.json -> 输出 .docx（按 skill 规定结构）。

problems.json 结构:
    {
      "title": "三年级 · 巧求面积 · 举一反三",
      "filename": "3年级_巧求面积_举一反三_练习题.docx",
      "examples": [
        {"stem":"……题目……", "figure":"fig_例1.png",   # figure 可省略
         "analysis":"……", "solution":"……",
         "practices":["练习题1……","练习题2……"]},
        ...
      ],
      "homework": ["作业一题面","作业二题面", ...],          # 共12道
      "answers":  ["练习1.1：……","作业一：……", ...]         # 练习+作业全部答案
    }
"""
import json, sys, subprocess


def _ensure(mod, pip):
    try:
        __import__(mod)
    except ImportError:
        subprocess.run([sys.executable, "-m", "pip", "install", "-q", pip], check=False)


_ensure("docx", "python-docx")
from docx import Document
from docx.shared import Inches

CN = ["一", "二", "三", "四", "五", "六", "七", "八", "九", "十", "十一", "十二",
      "十三", "十四", "十五"]


def main():
    data = json.load(open(sys.argv[1], encoding="utf-8"))
    doc = Document()
    doc.add_heading(data.get("title", "小学奥数练习题"), 0)

    for i, ex in enumerate(data["examples"], 1):
        doc.add_paragraph(f"【例{i}】")
        doc.add_paragraph(ex.get("stem", ""))
        if ex.get("figure"):
            try:
                doc.add_picture(ex["figure"], width=Inches(3.2))
            except Exception:
                doc.add_paragraph(f"（图片缺失：{ex['figure']}）")
        if ex.get("analysis"):
            doc.add_paragraph("分析：" + ex["analysis"])
        if ex.get("solution"):
            doc.add_paragraph("解：" + ex["solution"])
        if ex.get("practices"):
            doc.add_paragraph("练习")
            for j, p in enumerate(ex["practices"], 1):
                doc.add_paragraph(f"练习{j}：{p}")
        doc.add_paragraph("")

    if data.get("homework"):
        doc.add_heading("作业", level=1)
        for k, h in enumerate(data["homework"]):
            num = CN[k] if k < len(CN) else str(k + 1)
            doc.add_paragraph(f"作业{num}：{h}")

    if data.get("answers"):
        doc.add_heading("参考答案", level=1)
        for a in data["answers"]:
            doc.add_paragraph(a)

    fn = data.get("filename", "小学奥数练习题.docx")
    doc.save(fn)
    print(fn)


if __name__ == "__main__":
    main()
